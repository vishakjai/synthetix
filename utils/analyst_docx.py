from __future__ import annotations

import html
import io
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from docx import Document  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    Document = None  # type: ignore
    WD_TABLE_ALIGNMENT = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    OxmlElement = None  # type: ignore
    qn = None  # type: ignore
    Inches = None  # type: ignore
    Pt = None  # type: ignore
    RGBColor = None  # type: ignore


COLOR_PRIMARY = "1F3864"
COLOR_SECONDARY = "2E75B6"
COLOR_SECTION_ACCENT = "1F6B75"
COLOR_TEXT = "111827"
COLOR_MUTED = "6B7280"
COLOR_WHITE = "FFFFFF"
COLOR_ROW_ALT = "F3F4F6"
COLOR_CARD_BG = "D6E4F0"
COLOR_ACCENT_BG = "D6EEF2"
COLOR_RISK_HIGH_BG = "FEE2E2"
COLOR_RISK_MED_BG = "FEF3C7"
COLOR_RISK_LOW_BG = "DCFCE7"
COLOR_GATE_PASS_BG = "DCFCE7"
COLOR_GATE_WARN_BG = "FEF3C7"
COLOR_GATE_FAIL_BG = "FEE2E2"
TEMPLATE_ENV_VAR = "ANALYST_DOCX_TEMPLATE_PATH"
_STRICT_TEMPLATE_ACTIVE = False


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _clean(value: Any) -> str:
    return str(value or "").strip()


def _to_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        return default


def _extract_report_and_raw(payload: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    safe = _as_dict(payload)
    if _clean(safe.get("artifact_type")) == "analyst_report":
        return safe, _as_dict(safe.get("raw_artifacts"))
    report = _as_dict(safe.get("analyst_report_v2"))
    if _clean(report.get("artifact_type")) == "analyst_report":
        return report, _as_dict(safe.get("raw_artifacts"))
    return safe, _as_dict(safe.get("raw_artifacts"))


def _collect_report_data(report: dict[str, Any], raw_artifacts: dict[str, Any] | None = None) -> dict[str, Any]:
    metadata = _as_dict(report.get("metadata"))
    project = _as_dict(metadata.get("project"))
    ctx = _as_dict(metadata.get("context_reference"))
    brief = _as_dict(report.get("decision_brief"))
    glance = _as_dict(brief.get("at_a_glance"))
    inventory = _as_dict(glance.get("inventory_summary"))
    strategy = _as_dict(brief.get("recommended_strategy"))
    decisions = _as_dict(brief.get("decisions_required"))
    top_risks = _as_list(brief.get("top_risks"))
    next_steps = _as_list(brief.get("next_steps"))
    delivery = _as_dict(report.get("delivery_spec"))
    backlog = _as_list(_as_dict(delivery.get("backlog")).get("items"))
    testing = _as_dict(delivery.get("testing_and_evidence"))
    quality_gates = _as_list(testing.get("quality_gates"))
    qa_report = _as_dict(report.get("qa_report_v1"))
    if not qa_report:
        qa_report = _as_dict(_as_dict(raw_artifacts).get("qa_report_v1")) if isinstance(raw_artifacts, dict) else {}
    qa_summary = _as_dict(qa_report.get("summary"))
    qa_structural = _as_dict(qa_report.get("structural"))
    qa_semantic = _as_dict(qa_report.get("semantic"))
    qa_quality_gates = _as_list(qa_report.get("quality_gates"))
    open_questions = _as_list(delivery.get("open_questions"))
    appendix = _as_dict(report.get("appendix"))
    artifact_refs = _as_dict(appendix.get("artifact_refs"))
    raw = _as_dict(raw_artifacts)
    return {
        "metadata": metadata,
        "project": project,
        "context": ctx,
        "brief": brief,
        "glance": glance,
        "inventory": inventory,
        "strategy": strategy,
        "decisions": decisions,
        "top_risks": top_risks,
        "next_steps": next_steps,
        "delivery": delivery,
        "backlog": backlog,
        "testing": testing,
        "quality_gates": quality_gates,
        "qa_report": qa_report,
        "qa_summary": qa_summary,
        "qa_structural_checks": _as_list(qa_structural.get("checks")),
        "qa_semantic_checks": _as_list(qa_semantic.get("checks")),
        "qa_quality_gates": qa_quality_gates,
        "open_questions": open_questions,
        "artifact_refs": artifact_refs,
        "raw_artifacts": raw,
    }


def _set_run_style(run: Any, *, size: int = 11, bold: bool = False, color: str = COLOR_TEXT) -> None:
    if run is None or Pt is None or RGBColor is None:
        return
    run.bold = bool(bold)
    # In strict-template mode, keep the template's font family but still
    # apply emphasis/size/color so generated content remains visually rich.
    if not _STRICT_TEMPLATE_ACTIVE:
        run.font.name = "Arial"
    run.font.size = Pt(size)
    try:
        run.font.color.rgb = RGBColor.from_string(color)
    except Exception:
        pass


def _apply_heading_style(doc: Any) -> None:
    if doc is None or Pt is None or RGBColor is None:
        return
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
    except Exception:
        pass
    for style_name, size, color in [
        ("Title", 28, COLOR_PRIMARY),
        ("Heading 1", 16, COLOR_PRIMARY),
        ("Heading 2", 13, COLOR_SECONDARY),
        ("Heading 3", 12, COLOR_PRIMARY),
    ]:
        try:
            st = doc.styles[style_name]
            st.font.name = "Arial"
            st.font.bold = True
            st.font.size = Pt(size)
            st.font.color.rgb = RGBColor.from_string(color)
        except Exception:
            continue


def _set_cell_shading(cell: Any, fill_hex: str) -> None:
    if cell is None or OxmlElement is None or qn is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(fill_hex or "").replace("#", ""))
    tc_pr.append(shd)


def _set_cell_text(cell: Any, text: str, *, bold: bool = False, size: int = 10, color: str = COLOR_TEXT) -> None:
    if cell is None:
        return
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or ""))
    _set_run_style(run, size=size, bold=bold, color=color)


def _add_labeled_paragraph(doc: Any, label: str, value: str, *, muted: bool = False) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    _set_run_style(r1, size=10, bold=True, color=COLOR_SECONDARY)
    r2 = p.add_run(value or "n/a")
    _set_run_style(r2, size=10, bold=False, color=COLOR_MUTED if muted else COLOR_TEXT)


def _severity_fill(sev: str) -> str:
    s = str(sev or "").strip().lower()
    if s in {"high", "critical", "blocker"}:
        return COLOR_RISK_HIGH_BG
    if s in {"medium", "warn"}:
        return COLOR_RISK_MED_BG
    return COLOR_RISK_LOW_BG


def _gate_fill(result: str) -> str:
    s = str(result or "").strip().lower()
    if s == "pass":
        return COLOR_GATE_PASS_BG
    if s == "fail":
        return COLOR_GATE_FAIL_BG
    return COLOR_GATE_WARN_BG


def _base_form_name(value: Any) -> str:
    text = _clean(value)
    if not text:
        return ""
    if "::" in text:
        text = text.split("::")[-1]
    text = text.split(":")[0]
    text = text.rsplit("/", 1)[-1]
    if text.lower().endswith(".frm"):
        text = text[:-4]
    return text.lower()


def _extract_forms_from_text(value: Any) -> list[str]:
    text = _clean(value)
    if not text:
        return []
    forms = re.findall(r"([A-Za-z0-9_./-]+)", text)
    out: list[str] = []
    for form in forms:
        low = form.lower()
        if low.endswith((".bas", ".cls", ".ctl", ".vbp", ".vbg", ".res", ".dsr", ".mdl", ".mod")):
            continue
        if "frm" not in low and not low.startswith("form"):
            continue
        normalized = form.rsplit("/", 1)[-1]
        normalized = normalized[:-4] if normalized.lower().endswith(".frm") else normalized
        normalized = normalized.rstrip(".,;:()[]{}")
        nlow = normalized.lower()
        if re.search(r"_(click|change|load|keypress|gotfocus|lostfocus|activate|deactivate)$", nlow):
            continue
        if normalized not in out:
            out.append(normalized)
    return out


def _project_label(name: str, path_map: dict[str, str]) -> str:
    key = _clean(name)
    if not key:
        return "n/a"
    path = _clean(path_map.get(key))
    return f"{key} [{path}]" if path else key


def _project_from_scoped(value: Any) -> str:
    text = _clean(value)
    if "::" in text:
        return _clean(text.split("::", 1)[0])
    return ""


def _form_key(project_name: Any, form_name: Any) -> str:
    form = _base_form_name(form_name)
    project = _clean(project_name).lower()
    return f"{project}::{form}" if project and form else form


def _base_only_key(form_name: Any) -> str:
    base = _base_form_name(form_name)
    return f"__base__::{base}" if base else ""


def _form_keys(project_name: Any, form_name: Any) -> list[str]:
    keys: list[str] = []
    scoped = _form_key(project_name, form_name)
    if scoped:
        keys.append(scoped)
    base_key = _base_only_key(form_name)
    if base_key and base_key not in keys:
        keys.append(base_key)
    return keys


def _lookup_rows(mapping: dict[str, list[dict[str, Any]]], project_name: Any, form_name: Any) -> list[dict[str, Any]]:
    scoped = _form_key(project_name, form_name)
    base_key = _base_only_key(form_name)
    if scoped and scoped in mapping and mapping[scoped]:
        return mapping[scoped]
    if base_key and base_key in mapping and mapping[base_key]:
        return mapping[base_key]
    return []


def _lookup_set(mapping: dict[str, set[str]], project_name: Any, form_name: Any) -> set[str]:
    scoped = _form_key(project_name, form_name)
    base_key = _base_only_key(form_name)
    if scoped and scoped in mapping and mapping[scoped]:
        return mapping[scoped]
    if base_key and base_key in mapping and mapping[base_key]:
        return mapping[base_key]
    return set()


def _lookup_control_map(mapping: dict[str, dict[str, str]], project_name: Any, form_name: Any) -> dict[str, str]:
    scoped = _form_key(project_name, form_name)
    base_key = _base_only_key(form_name)
    if scoped and scoped in mapping and mapping[scoped]:
        return mapping[scoped]
    if base_key and base_key in mapping and mapping[base_key]:
        return mapping[base_key]
    return {}


def _qualified_form_name(project_name: Any, form_name: Any) -> str:
    project = _clean(project_name)
    form = _clean(form_name)
    if project and form:
        return f"{project}::{form}"
    return form or "n/a"


def _split_words(token: str) -> str:
    raw = _clean(token)
    lowered = raw.lower()
    if lowered.startswith("dtpicker"):
        raw = raw[len("dtpicker"):]
        raw = f"date{raw}" if raw else "date"
        lowered = raw.lower()
    for prefix in ("txt", "cbo", "cmb", "dtp", "msk", "lst", "chk", "opt", "lbl", "cmd"):
        if lowered.startswith(prefix) and len(raw) > len(prefix):
            lower_raw = raw.lower()
            if prefix == "opt" and lower_raw.startswith("option"):
                continue
            if prefix == "chk" and lower_raw.startswith("check"):
                continue
            if prefix == "txt" and lower_raw.startswith("text"):
                continue
            if prefix == "cbo" and lower_raw.startswith("combo"):
                continue
            raw = raw[len(prefix):]
            break
    raw = re.sub(r"[_\\-]+", " ", raw)
    raw = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", raw)
    raw = re.sub(r"([A-Za-z])(id|no)\b", r"\1 \2", raw, flags=re.IGNORECASE)
    raw = re.sub(r"([A-Za-z])([0-9])", r"\1 \2", raw)
    raw = re.sub(r"\s+", " ", raw).strip()
    return raw.lower()


def _is_data_input_control(control_id: str) -> bool:
    cid = _clean(control_id).lower()
    return cid.startswith(("txt", "cbo", "cmb", "dtp", "msk", "lst", "chk", "opt"))


def _to_business_input(control_id: str) -> str:
    words = _split_words(control_id)
    return words or _clean(control_id).lower()


def _callable_kind(procedure_name: Any, form_name: Any, event_hint: Any = "") -> str:
    proc = _clean(procedure_name).lower()
    form = _base_form_name(form_name)
    evt = _clean(event_hint).lower()
    if form == "shared_module":
        return "shared_function"
    if evt:
        return "event_handler"
    if re.search(r"_(click|change|load|keypress|keydown|keyup|gotfocus|lostfocus|activate|deactivate)$", proc):
        return "event_handler"
    if proc.startswith(("cmd", "lbl", "txt", "cbo", "opt", "chk")):
        return "event_handler"
    return "procedure"


def _semantic_form_alias(
    *,
    form_name: str,
    purpose: str,
    db_tables: set[str],
    procedures: list[dict[str, Any]],
    rules: list[dict[str, Any]],
    controls: list[str] | None = None,
) -> str:
    form_token = _clean(form_name).lower()
    if form_token in {"main", "mdiform"}:
        return "Navigation Hub"
    is_generic_form = bool(re.fullmatch(r"(form\d+|frm\d+)", form_token))
    purpose_low = _clean(purpose).lower()
    if "deposit capture" in purpose_low:
        return "Deposit Capture"
    if "withdrawal processing" in purpose_low:
        return "Withdrawal Processing"
    if "customer profile" in purpose_low:
        return "Customer Management"
    if "transaction ledger" in purpose_low:
        return "Transaction Ledger"
    if "account type maintenance" in purpose_low:
        return "Account Type Maintenance"
    token_blob = " ".join(
        [
            form_token,
            _clean(purpose).lower(),
            " ".join(_clean(x).lower() for x in db_tables if _clean(x)),
            " ".join(_clean(_as_dict(p).get("procedure_name")).lower() for p in procedures),
            " ".join(_clean(_as_dict(r).get("statement")).lower() for r in rules),
            " ".join(_clean(x).lower() for x in (controls or []) if _clean(x)),
        ]
    )
    if any(key in token_blob for key in ("login", "logi", "username", "password", "txtpass", "pass1", "credential")):
        return "Password Management" if any(key in token_blob for key in ("txtpass", "pass1", "credential")) else "Authentication"
    has_strong_transaction_signal = any(
        key in token_blob for key in ("transction", "tbltransaction", "transaction ledger", "ledger")
    )
    if has_strong_transaction_signal or ("debit" in token_blob and "credit" in token_blob):
        return "Transaction Ledger"
    if any(key in token_blob for key in ("withdraw", "debit")):
        return "Withdrawal Processing"
    if any(key in token_blob for key in ("deposit", "credit", "balancedt")) and not any(
        key in token_blob for key in ("transaction", "transction", "debit")
    ):
        return "Deposit Capture"
    if any(key in token_blob for key in ("customer", "tblcustomer")) and any(
        key in token_blob for key in ("interest", "min balance", "account type", "acctype")
    ):
        return "Customer Management"
    if any(key in token_blob for key in ("accounttype", "acctype")):
        return "Account Type Maintenance"
    if any(key in token_blob for key in ("customer", "tblcustomer")):
        return "Customer Management"
    if any(key in token_blob for key in ("report", "datareport", "dataenvironment")):
        return "Reporting"
    if any(key in token_blob for key in ("search", "lookup", "find")):
        return "Search"
    if any(key in token_blob for key in ("main", "mdiform", "toolbar")):
        return "Navigation Hub"
    if any(key in token_blob for key in ("balance", "tblbalance")):
        return "Balance Inquiry"
    if any(key in token_blob for key in ("timer", "progressbar", "splash")):
        return "Splash/Loading"
    if is_generic_form and form_token == "form9":
        return "Authentication Entry"
    if is_generic_form and form_token == "form1":
        return "Navigation/Menu"
    if is_generic_form and any(key in token_blob for key in ("dated", "datejoined", "dtpicker", "date 1", "from date", "to date")):
        return "Date/Period Entry"
    cleaned_purpose = _clean(purpose).rstrip(".")
    if cleaned_purpose:
        short = re.sub(r"\bworkflow\b", "", cleaned_purpose, flags=re.IGNORECASE)
        short = re.sub(r"\s+", " ", short).strip(" -")
        generic_phrases = {
            "business executed through event-driven ui controls",
            "business workflow executed through event-driven ui controls",
            "application navigation and module routing",
        }
        if short.lower() in generic_phrases:
            return ""
        if short:
            return short
    return ""


def _display_form_name(form_name: str, alias: str) -> str:
    generic = bool(re.fullmatch(r"(form\d+|frm\d+)", _clean(form_name).lower()))
    if generic and _clean(alias):
        return f"{_clean(form_name)} [{_clean(alias)}]"
    return _clean(form_name)


def _rule_business_meaning(statement: str, category: str) -> str:
    stmt = _clean(statement)
    low = stmt.lower()
    if (
        ("asc(" in low and ("< 46" in low or "<= 45" in low) and ("> 57" in low or ">= 58" in low))
        or (
            ("keyascii" in low or "keyvalue" in low)
            and (">= 48" in low and "<= 57" in low)
        )
    ):
        return "Input is restricted to numeric digits only."
    if re.search(r"keyascii\s*=\s*13", low):
        return "Pressing Enter triggers the same action flow as the primary button."
    if "case keyascii" in low:
        return "Keyboard input routing determines which action path is executed."
    if re.search(r"\.state\s*=\s*1", low):
        return "The action proceeds only when the recordset/connection is active."
    if re.search(r"\.recordcount\s*>\s*0", low):
        return "The action proceeds only when matching records are found."
    if ("max(" in low and "+ 1" in low) or ("max(" in low and "+1" in low):
        return "A new identifier is generated as current maximum value plus one."
    if "computed value rule" in low or "currbalance" in low:
        if "lblbalance.caption" in low:
            return "Balance is recalculated from the displayed balance label and entered amount (UI-derived source)."
        return "Balance is recalculated using the entered amount and current account value."
    if "case button.index" in low or "case buttonmenu.key" in low:
        return "User menu selection routes the workflow to the corresponding module."
    if "threshold decision rule" in low and "if " in low:
        cond = stmt.split("IF", 1)[-1].strip()
        cond = re.sub(r"\s*THEN.*$", "", cond, flags=re.IGNORECASE).strip()
        cond_low = cond.lower()
        if (
            ("asc(" in cond_low and ("< 46" in cond_low or "<= 45" in cond_low) and ("> 57" in cond_low or ">= 58" in cond_low))
            or (
                ("keyascii" in cond_low or "keyvalue" in cond_low)
                and (">= 48" in cond_low and "<= 57" in cond_low)
            )
        ):
            return "Input is restricted to numeric digits only."
        return f"The workflow continues only when this condition is true: {cond}."
    if "executes transaction workflow through procedures" in low:
        return "Workflow is orchestrated through UI event handlers and internal procedures."
    if "reads/writes persisted entities" in low:
        return "Form persists and retrieves records from the listed tables."
    if "authenticate users" in low:
        return "User authentication is required before entering the workflow."
    if category.lower() in {"data_persistence", "calculation_logic", "threshold_rule"}:
        return "Business behavior is enforced through data and calculation logic."
    return stmt


def _business_effect_from_sql(op: str, table: str) -> str:
    op_low = _clean(op).lower()
    table_low = _clean(table).lower()
    if op_low == "select":
        if "balance" in table_low:
            return "Customer balance and account details displayed for review."
        if "customer" in table_low:
            return "Customer details displayed for review."
        if "transction" in table_low or "transaction" in table_low:
            return "Transaction history displayed for review."
        if "accounttype" in table_low:
            return "Account type details displayed for selection."
        if "logi" in table_low or "login" in table_low:
            return "User credentials validated against stored records."
    if "deposit" in table_low:
        return "Deposit transaction recorded."
    if "withdraw" in table_low:
        return "Withdrawal transaction recorded."
    if "balance" in table_low:
        return "Account balance updated."
    if "transction" in table_low or "transaction" in table_low:
        return "Transaction ledger updated."
    if "customer" in table_low and op_low in {"insert", "update", "delete"}:
        return "Customer profile data updated."
    if "accounttype" in table_low:
        return "Account type configuration updated."
    if "logi" in table_low or "login" in table_low:
        return "User authentication record validated."
    if op_low == "insert":
        return f"New record created in {table}."
    if op_low == "update":
        return f"Existing records updated in {table}."
    if op_low == "delete":
        return f"Records deleted from {table}."
    return ""


def _fallback_business_effects(
    *,
    alias: str,
    purpose: str,
    inputs: set[str],
    db_tables: list[str],
    rules: list[dict[str, Any]],
) -> list[str]:
    token_blob = " ".join(
        [
            _clean(alias).lower(),
            _clean(purpose).lower(),
            " ".join(_clean(x).lower() for x in sorted(inputs)),
            " ".join(_clean(x).lower() for x in db_tables),
            " ".join(_clean(_as_dict(r).get("statement")).lower() for r in rules),
        ]
    )
    effects: list[str] = []
    if any(x in token_blob for x in ["deposit", "amount deposited", "credit"]):
        effects.append("Deposit transaction recorded.")
        effects.append("Account balance recalculated.")
    if any(x in token_blob for x in ["withdraw", "amount withdrawn", "debit"]):
        effects.append("Withdrawal transaction recorded.")
        effects.append("Account balance recalculated.")
    if any(x in token_blob for x in ["transaction ledger", "transction", "transaction"]):
        effects.append("Transaction history updated.")
    if any(x in token_blob for x in ["customer management", "customer profile"]):
        effects.append("Customer profile created or updated.")
    if any(x in token_blob for x in ["account type", "acctype", "accounttype"]):
        effects.append("Account type master data maintained.")
    if any(x in token_blob for x in ["authentication", "login", "password"]):
        effects.append("User access is validated before workflow continuation.")
    if any(x in token_blob for x in ["search", "lookup"]):
        effects.append("Matching records displayed to the user.")
    if any(x in token_blob for x in ["navigation hub", "main", "toolbar"]):
        effects.append("Navigation routes the user to selected module screens.")
    if not effects and db_tables:
        if any("balance" in _clean(t).lower() for t in db_tables):
            effects.append("Account balance information refreshed.")
        if any("customer" in _clean(t).lower() for t in db_tables):
            effects.append("Customer details loaded/updated for the selected workflow.")
    dedup: list[str] = []
    for eff in effects:
        e = _clean(eff)
        if e and e not in dedup:
            dedup.append(e)
    return dedup[:6]


def _infer_form_type(*, form_name: str, purpose: str, procedures: list[dict[str, Any]], controls: list[Any], tables: set[str]) -> str:
    form_low = form_name.lower()
    purpose_low = purpose.lower()
    control_text = " ".join(_clean(c).lower() for c in controls)
    proc_names = {_clean(_as_dict(p).get("procedure_name")).lower() for p in procedures}
    if "splash" in form_low or "splash" in purpose_low:
        return "Splash"
    if (
        form_low in {"main", "mdiform"}
        or form_low.startswith("mdi")
        or "toolbar" in control_text
        or any("toolbar" in x for x in proc_names)
    ):
        return "MDI_Host"
    if (
        "login" in form_low
        or "auth" in purpose_low
        or ("form9" in form_low and ({"logi", "login"} & {t.lower() for t in tables}))
    ):
        return "Login"
    if form_low.startswith(("rpt", "datareport")):
        return "Report"
    return "Child"


def _normalize_data_touchpoints(values: list[str]) -> list[str]:
    known_tables = {
        "accounttype",
        "balancedt",
        "customer",
        "deposit",
        "withdrawal",
        "transactions",
        "transctions",
        "login",
        "logi",
        "tblbalances",
        "tblcustomers",
        "tbltransactions",
    }
    out: list[str] = []
    for value in values:
        token = _clean(value).lower()
        if not token:
            continue
        if token in {"con", "rs", "label", "button", "textbox", "combobox"}:
            continue
        if token.startswith(("cmd", "txt", "lbl", "cbo", "cmb", "chk", "opt", "lst", "frm", "form", "module")):
            continue
        if token in {"accountno", "accountid", "customerid", "transactionid", "dated", "datejoined", "address", "amount"}:
            continue
        if token in known_tables or token.startswith("tbl"):
            if token not in out:
                out.append(token)
    return out[:16]


def _resolve_paragraph_style(doc: Any, role: str) -> Any:
    aliases = {
        "title": ["Title"],
        "heading1": ["Heading 1"],
        "heading2": ["Heading 2"],
        "heading3": ["Heading 3"],
        "list": ["List Paragraph"],
        "body": ["Normal"],
    }
    role_key = str(role or "").strip().lower()
    names = aliases.get(role_key, [])
    paragraph_styles = [
        style
        for style in getattr(doc, "styles", [])
        if str(getattr(style, "type", "")).endswith("PARAGRAPH (1)")
    ]
    for name in names:
        for style in paragraph_styles:
            if _clean(getattr(style, "name", "")) == name:
                return style
    for style in paragraph_styles:
        name = _clean(getattr(style, "name", ""))
        style_id = _clean(getattr(style, "style_id", ""))
        if role_key == "heading1" and style_id.lower() == "heading1":
            return style
        if role_key == "heading2" and style_id.lower() == "heading2":
            return style
        if role_key == "list" and style_id.lower() == "listparagraph":
            return style
        if role_key == "title" and style_id.lower() == "title":
            return style
        if role_key == "body" and (name.lower() == "normal" or style_id.lower() == "normal"):
            return style
    return None


def _add_paragraph_with_role(doc: Any, text: str = "", *, role: str = "body") -> Any:
    style_obj = _resolve_paragraph_style(doc, role)
    if style_obj is not None:
        return doc.add_paragraph(text, style=style_obj)
    return doc.add_paragraph(text)


def _resolve_table_style(doc: Any) -> Any:
    preferred = ["Table Grid"]
    for name in preferred:
        try:
            return doc.styles[name]
        except Exception:
            continue
    for style in getattr(doc, "styles", []):
        if str(getattr(style, "type", "")).endswith("TABLE (3)"):
            return style
    return None


def _build_table(doc: Any, headers: list[str], *, col_count: int | None = None) -> Any:
    cols = int(col_count or len(headers) or 1)
    table = doc.add_table(rows=1, cols=cols)
    resolved_style = _resolve_table_style(doc)
    if resolved_style is not None:
        try:
            table.style = resolved_style
        except Exception:
            pass
    if WD_TABLE_ALIGNMENT:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, head in enumerate(headers[:cols]):
        cell = table.rows[0].cells[idx]
        _set_cell_shading(cell, COLOR_PRIMARY)
        _set_cell_text(cell, head, bold=True, color=COLOR_WHITE, size=9)
    return table


def _add_table_row(table: Any, values: list[Any], *, alt: bool = False, size: int = 9) -> None:
    row_cells = table.add_row().cells
    for idx, value in enumerate(values[: len(row_cells)]):
        text = _clean(value) or "n/a"
        _set_cell_text(row_cells[idx], text, size=size, color=COLOR_TEXT)
        if alt:
            _set_cell_shading(row_cells[idx], COLOR_ROW_ALT)


def _normalize_llm_doc_plan(value: Any) -> dict[str, Any]:
    plan = _as_dict(value)
    section_intros = _as_dict(plan.get("section_intros"))
    callouts = _as_list(plan.get("callouts"))
    bullets = _as_list(plan.get("executive_bullets"))
    return {
        "title": _clean(plan.get("title"))[:140],
        "subtitle": _clean(plan.get("subtitle"))[:220],
        "narrative": _clean(plan.get("narrative"))[:2200],
        "executive_bullets": [_clean(x)[:260] for x in bullets if _clean(x)][:8],
        "callouts": [
            {
                "label": _clean(_as_dict(x).get("label"))[:64],
                "message": _clean(_as_dict(x).get("message"))[:260],
                "severity": _clean(_as_dict(x).get("severity") or "info").lower(),
            }
            for x in callouts[:8]
            if isinstance(x, dict) and (_clean(_as_dict(x).get("label")) or _clean(_as_dict(x).get("message")))
        ],
        "section_intros": {
            "executive_snapshot": _clean(section_intros.get("executive_snapshot"))[:320],
            "dependency_map": _clean(section_intros.get("dependency_map"))[:320],
            "form_dossiers": _clean(section_intros.get("form_dossiers"))[:320],
            "flow_traces": _clean(section_intros.get("flow_traces"))[:320],
            "traceability": _clean(section_intros.get("traceability"))[:320],
            "sprints": _clean(section_intros.get("sprints"))[:320],
            "risks": _clean(section_intros.get("risks"))[:320],
        },
    }


def _add_section_intro(doc: Any, text: str) -> None:
    if not _clean(text):
        return
    para = doc.add_paragraph()
    run = para.add_run(_clean(text))
    _set_run_style(run, size=9, color=COLOR_MUTED)


def _resolve_docx_template_path(template_path: str | None = None) -> str:
    requested = _clean(template_path)
    if requested and Path(requested).exists():
        return requested

    env_path = _clean(os.getenv(TEMPLATE_ENV_VAR, ""))
    if env_path and Path(env_path).exists():
        return env_path

    module_root = Path(__file__).resolve().parents[1]
    candidates = [
        module_root / "assets" / "docx_templates" / "ba_workbook_v3.docx",
        Path("/Users/vishak/Downloads/ba_workbook_v3.docx"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return ""


def _clear_document_body(doc: Any) -> None:
    body = getattr(getattr(doc, "_element", None), "body", None)
    if body is None:
        return
    children = list(body)
    for child in children:
        # Keep final section properties so page config/header/footer links remain valid.
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _build_business_docx_bytes_rich(
    payload: dict[str, Any],
    *,
    run_id: str = "",
    llm_doc_plan: dict[str, Any] | None = None,
    template_path: str | None = None,
    strict_template: bool = False,
) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not available")
    global _STRICT_TEMPLATE_ACTIVE
    _STRICT_TEMPLATE_ACTIVE = bool(strict_template)
    report, raw_artifacts = _extract_report_and_raw(payload)
    resolved_template = _resolve_docx_template_path(template_path) if strict_template else _clean(template_path)
    if resolved_template:
        doc = Document(resolved_template)
        _clear_document_body(doc)
    else:
        doc = Document()
    if not strict_template:
        _apply_heading_style(doc)
    data = _collect_report_data(report, raw_artifacts)
    raw = _as_dict(data.get("raw_artifacts"))

    section = doc.sections[0]
    if Inches is not None and not strict_template:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)

    if not strict_template:
        header = section.header.paragraphs[0]
        header.clear()
        hrun = header.add_run("VB6 Legacy Analyst Workbook")
        _set_run_style(hrun, size=10, bold=True, color=COLOR_SECTION_ACCENT)
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT if WD_ALIGN_PARAGRAPH else 0

        footer = section.footer.paragraphs[0]
        footer.clear()
        frun = footer.add_run("Synthetix Platform | Analyst Evidence Workbook")
        _set_run_style(frun, size=9, bold=False, color=COLOR_MUTED)
        footer.alignment = WD_ALIGN_PARAGRAPH.LEFT if WD_ALIGN_PARAGRAPH else 0

    banner = _build_table(doc, ["VB6 Modernization - Business Analyst Workbook"], col_count=1)
    _set_cell_shading(banner.rows[0].cells[0], COLOR_PRIMARY)
    plan = _normalize_llm_doc_plan(llm_doc_plan)
    workbook_title = plan.get("title") or "VB6 Modernization - Business Analyst Workbook"
    _set_cell_text(banner.rows[0].cells[0], workbook_title, bold=True, color=COLOR_WHITE, size=12)

    project_name = _clean(data["project"].get("name")) or "Untitled Project"
    subtitle = doc.add_paragraph()
    _set_run_style(subtitle.add_run(project_name), size=13, bold=True, color=COLOR_SECTION_ACCENT)
    if _clean(plan.get("subtitle")):
        subtitle2 = doc.add_paragraph()
        _set_run_style(subtitle2.add_run(_clean(plan.get("subtitle"))), size=10, color=COLOR_TEXT)
    stamp = _clean(data["metadata"].get("generated_at")) or datetime.now(timezone.utc).isoformat()
    meta = doc.add_paragraph()
    _set_run_style(
        meta.add_run(
            f"Run {run_id or 'n/a'}  ·  Generated {stamp}  ·  Repo {_clean(data['context'].get('repo')) or 'n/a'}"
        ),
        size=9,
        color=COLOR_MUTED,
    )

    doc.add_paragraph()
    _add_paragraph_with_role(doc, "Executive Snapshot", role="heading1")
    _add_section_intro(doc, _as_dict(plan.get("section_intros")).get("executive_snapshot", ""))
    objective = _clean(data["project"].get("objective"))
    if objective:
        _add_labeled_paragraph(doc, "Objective", objective)
    headline = _clean(data["glance"].get("headline"))
    if headline:
        _add_labeled_paragraph(doc, "Summary", headline)
    narrative = _clean(plan.get("narrative"))
    if narrative:
        n = doc.add_paragraph()
        _set_run_style(n.add_run(narrative), size=10, color=COLOR_TEXT)
    bullets = _as_list(plan.get("executive_bullets"))
    for bullet in bullets:
        bp = _add_paragraph_with_role(doc, role="list")
        _set_run_style(bp.add_run(f"- {_clean(bullet)}"), size=9, color=COLOR_TEXT)

    inv = data["inventory"]
    tables_touched_raw = [str(x).strip() for x in _as_list(inv.get("tables_touched")) if str(x).strip()]
    tables_touched = _normalize_data_touchpoints(tables_touched_raw)
    summary_table = _build_table(doc, ["Category", "Summary"])
    summary_rows = [
        ("Readiness", f"{_clean(data['glance'].get('readiness_score')) or 'n/a'}/100"),
        ("Risk tier", _clean(data["glance"].get("risk_tier")) or "n/a"),
        (
            "Inventory",
            f"projects={_to_int(inv.get('projects'))}, forms={_to_int(inv.get('forms'))}, dependencies={_to_int(inv.get('dependencies'))}",
        ),
        ("Data touchpoints", ", ".join(tables_touched[:16]) if tables_touched else "n/a"),
    ]
    for idx, row in enumerate(summary_rows):
        _add_table_row(summary_table, [row[0], row[1]], alt=(idx % 2 == 1))
    for idx, callout in enumerate(_as_list(plan.get("callouts"))[:4]):
        row = _as_dict(callout)
        label = _clean(row.get("label")) or f"Callout {idx + 1}"
        message = _clean(row.get("message")) or "n/a"
        sev = _clean(row.get("severity") or "info")
        c = doc.add_paragraph()
        _set_run_style(c.add_run(f"{label}: "), size=9, bold=True, color=COLOR_SECONDARY)
        _set_run_style(c.add_run(message), size=9, color=COLOR_TEXT)
        if sev in {"high", "critical", "blocker"}:
            c_format = c.paragraph_format
            c_format.left_indent = Inches(0.08) if Inches else None

    strategy = data["strategy"]
    _add_paragraph_with_role(doc, "Recommended Modernization Strategy", role="heading2")
    p = doc.add_paragraph()
    _set_run_style(p.add_run(f"{_clean(strategy.get('name')) or 'Phased migration'}: "), size=10, bold=True, color=COLOR_SECTION_ACCENT)
    _set_run_style(p.add_run(_clean(strategy.get("rationale")) or "Rationale not provided."), size=10, color=COLOR_TEXT)
    for phase in _as_list(strategy.get("phases"))[:8]:
        if not isinstance(phase, dict):
            continue
        item = _add_paragraph_with_role(doc, role="list")
        _set_run_style(item.add_run(f"- {_clean(phase.get('id'))} {_clean(phase.get('title'))}: "), size=9, bold=True, color=COLOR_SECONDARY)
        _set_run_style(item.add_run(_clean(phase.get("outcome"))), size=9, color=COLOR_TEXT)

    decisions = _as_dict(data.get("decisions"))
    decision_rows = _as_list(decisions.get("blocking"))[:12] + _as_list(decisions.get("non_blocking"))[:8]
    _add_paragraph_with_role(doc, "Decisions Required", role="heading2")
    dec_table = _build_table(doc, ["ID", "Question", "Recommendation"])
    if not decision_rows:
        decision_rows = [{"id": "DEC-NA", "question": "No explicit decisions listed.", "default_recommendation": ""}]
    for idx, row in enumerate(decision_rows):
        r = _as_dict(row)
        _add_table_row(
            dec_table,
            [
                _clean(r.get("id")) or "DEC",
                _clean(r.get("question")) or "n/a",
                _clean(r.get("default_recommendation")) or "n/a",
            ],
            alt=(idx % 2 == 1),
        )

    quality_gates = _as_list(data.get("quality_gates"))
    _add_paragraph_with_role(doc, "Quality Gates", role="heading2")
    gate_table = _build_table(doc, ["Gate ID", "Result", "Description"])
    if not quality_gates:
        quality_gates = [{"id": "gate_na", "result": "warn", "description": "No quality gates listed."}]
    for idx, row in enumerate(quality_gates[:24]):
        r = _as_dict(row)
        result = _clean(r.get("result")) or "warn"
        _add_table_row(
            gate_table,
            [
                _clean(r.get("id")) or "gate",
                result.upper(),
                _clean(r.get("description")) or "n/a",
            ],
            alt=(idx % 2 == 1),
        )
        _set_cell_shading(gate_table.rows[idx + 1].cells[1], _gate_fill(result))

    qa_summary = _as_dict(data.get("qa_summary"))
    qa_structural_checks = _as_list(data.get("qa_structural_checks"))
    qa_semantic_checks = _as_list(data.get("qa_semantic_checks"))
    qa_quality_gates = _as_list(data.get("qa_quality_gates"))
    _add_paragraph_with_role(doc, "QA Summary", role="heading2")
    qa_overview = _build_table(
        doc,
        ["Status", "Pass", "Warn", "Fail", "Blockers", "Semantic Warn"],
    )
    _add_table_row(
        qa_overview,
        [
            _clean(qa_summary.get("status")) or "NOT_RUN",
            str(_to_int(qa_summary.get("pass_count"))),
            str(_to_int(qa_summary.get("warn_count"))),
            str(_to_int(qa_summary.get("fail_count"))),
            str(_to_int(qa_summary.get("blocker_count"))),
            str(_to_int(qa_summary.get("semantic_warn_count"))),
        ],
        alt=False,
    )
    if qa_quality_gates:
        qa_gate_table = _build_table(doc, ["QA Gate", "Result", "Description"])
        for idx, row in enumerate(qa_quality_gates[:20]):
            r = _as_dict(row)
            result = _clean(r.get("result")) or "warn"
            _add_table_row(
                qa_gate_table,
                [
                    _clean(r.get("id")) or "qa_gate",
                    result.upper(),
                    _clean(r.get("description")) or "n/a",
                ],
                alt=(idx % 2 == 1),
            )
            _set_cell_shading(qa_gate_table.rows[idx + 1].cells[1], _gate_fill(result))
    qa_structural_table = _build_table(doc, ["Check ID", "Result", "Blocking", "Detail"])
    if qa_structural_checks:
        for idx, row in enumerate(qa_structural_checks[:40]):
            r = _as_dict(row)
            result = _clean(r.get("result")) or "warn"
            _add_table_row(
                qa_structural_table,
                [
                    _clean(r.get("check_id") or r.get("id")) or f"check_{idx+1}",
                    result.upper(),
                    "yes" if bool(r.get("blocking")) else "no",
                    _clean(r.get("detail")) or "n/a",
                ],
                alt=(idx % 2 == 1),
            )
            _set_cell_shading(qa_structural_table.rows[idx + 1].cells[1], _gate_fill(result))
    else:
        _add_table_row(qa_structural_table, ["none", "PASS", "no", "No structural checks emitted."], alt=False)
    qa_semantic_table = _build_table(doc, ["Check ID", "Severity", "Confidence", "Detail"])
    if qa_semantic_checks:
        for idx, row in enumerate(qa_semantic_checks[:40]):
            r = _as_dict(row)
            severity = _clean(r.get("severity")) or "medium"
            _add_table_row(
                qa_semantic_table,
                [
                    _clean(r.get("check_id") or r.get("id")) or f"semantic_{idx+1}",
                    severity.upper(),
                    _clean(r.get("confidence")) or "n/a",
                    _clean(r.get("detail")) or "n/a",
                ],
                alt=(idx % 2 == 1),
            )
            _set_cell_shading(qa_semantic_table.rows[idx + 1].cells[1], _severity_fill(severity))
    else:
        _add_table_row(qa_semantic_table, ["none", "LOW", "n/a", "No semantic warnings."], alt=False)

    # Raw artifacts for detailed BA sections
    raw_legacy = _as_dict(raw.get("legacy_inventory"))
    projects = _as_list(raw_legacy.get("projects"))
    raw_event_map = _as_list(_as_dict(raw.get("event_map")).get("entries"))
    raw_sql_map = _as_list(_as_dict(raw.get("sql_map")).get("entries"))
    raw_procedures = _as_list(_as_dict(raw.get("procedure_summary")).get("procedures"))
    raw_form_dossiers = _as_list(_as_dict(raw.get("form_dossier")).get("dossiers"))
    raw_rules = _as_list(_as_dict(raw.get("business_rule_catalog")).get("rules"))
    raw_risks = _as_list(_as_dict(raw.get("risk_register")).get("risks"))
    raw_orphans = _as_list(_as_dict(raw.get("orphan_analysis")).get("orphans"))
    raw_landscape = _as_list(_as_dict(raw.get("repo_landscape")).get("projects"))
    raw_variant_diff = _as_dict(raw.get("variant_diff_report"))

    project_path_by_name: dict[str, str] = {}
    project_dependencies_by_name: dict[str, set[str]] = {}
    project_tables_by_name: dict[str, set[str]] = {}
    project_display_name_counts: dict[str, int] = {}
    for row in raw_landscape:
        r = _as_dict(row)
        raw_id = _clean(r.get("id"))
        left = raw_id.split("|", 1)[0] if "|" in raw_id else raw_id
        project_display_name_counts[left or "variant"] = project_display_name_counts.get(left or "variant", 0) + 1
        path = _clean(r.get("path"))
        deps = set(_as_list(r.get("dependencies")))
        tables = set(_clean(t) for t in _as_list(r.get("db_touchpoints")) if _clean(t))
        names = [left, _clean(r.get("name")), raw_id]
        for name in names:
            key = _clean(name)
            if not key:
                continue
            if key not in project_path_by_name and path:
                project_path_by_name[key] = path
            project_dependencies_by_name.setdefault(key, set()).update(dep for dep in deps if _clean(dep))
            project_tables_by_name.setdefault(key, set()).update(tables)

    form_sql_rows: dict[str, list[dict[str, Any]]] = {}
    form_db_tables: dict[str, set[str]] = {}
    for row in raw_sql_map:
        r = _as_dict(row)
        tables = set(_clean(t) for t in _as_list(r.get("tables")) if _clean(t))
        project_name = _clean(r.get("variant")) or _project_from_scoped(r.get("form"))
        form_name = _clean(r.get("form_base")) or _clean(r.get("form"))
        keys = _form_keys(project_name, form_name)
        if not keys:
            continue
        for key in keys:
            form_sql_rows.setdefault(key, []).append(r)
            form_db_tables.setdefault(key, set()).update(tables)
        if project_name:
            project_tables_by_name.setdefault(project_name, set()).update(tables)

    form_event_rows: dict[str, list[dict[str, Any]]] = {}
    event_handler_by_key_proc: dict[tuple[str, str], set[str]] = {}
    form_shared_components: dict[str, set[str]] = {}
    for row in raw_event_map:
        r = _as_dict(row)
        container = _clean(r.get("container")) or _clean(r.get("name"))
        project_name = _project_from_scoped(container) or _project_from_scoped(_as_dict(r.get("handler")).get("symbol"))
        form_name = _base_form_name(container) or _base_form_name(_as_dict(r.get("handler")).get("symbol"))
        keys = _form_keys(project_name, form_name)
        if not keys:
            continue
        for key in keys:
            form_event_rows.setdefault(key, []).append(r)
        for call in [_clean(c) for c in _as_list(r.get("calls")) if _clean(c)]:
            for key in keys:
                event_handler_by_key_proc.setdefault((key, call), set()).add(
                    _clean(_as_dict(r.get("handler")).get("symbol")) or _clean(r.get("entry_id"))
                )

    form_proc_rows: dict[str, list[dict[str, Any]]] = {}
    for row in raw_procedures:
        r = _as_dict(row)
        project_name = _project_from_scoped(r.get("form"))
        form_name = _base_form_name(r.get("form"))
        for key in _form_keys(project_name, form_name):
            form_proc_rows.setdefault(key, []).append(r)

    form_risk_rows: dict[str, list[dict[str, Any]]] = {}
    for row in raw_risks:
        r = _as_dict(row)
        texts = [_clean(r.get("description")), _clean(r.get("recommended_action"))]
        for ev in _as_list(r.get("evidence")):
            e = _as_dict(ev)
            texts.append(_clean(_as_dict(e.get("external_ref")).get("ref")))
            texts.append(_clean(_as_dict(e.get("file_span")).get("path")))
        forms: set[str] = set()
        for text in texts:
            for form_name in _extract_forms_from_text(text):
                forms.add(_base_form_name(form_name))
        for dossier in raw_form_dossiers:
            d = _as_dict(dossier)
            if _base_form_name(d.get("form_name")) not in forms:
                continue
            for key in _form_keys(d.get("project_name"), d.get("form_name")):
                form_risk_rows.setdefault(key, []).append(r)

    form_rule_rows: dict[str, list[dict[str, Any]]] = {}
    for row in raw_rules:
        r = _as_dict(row)
        texts = [_clean(_as_dict(r.get("scope")).get("component_id")), _clean(r.get("statement"))]
        for ev in _as_list(r.get("evidence")):
            e = _as_dict(ev)
            texts.append(_clean(_as_dict(e.get("external_ref")).get("ref")))
            texts.append(_clean(_as_dict(e.get("file_span")).get("path")))
        forms: set[str] = set()
        for text in texts:
            for form_name in _extract_forms_from_text(text):
                forms.add(_base_form_name(form_name))
        for dossier in raw_form_dossiers:
            d = _as_dict(dossier)
            if _base_form_name(d.get("form_name")) not in forms:
                continue
            for key in _form_keys(d.get("project_name"), d.get("form_name")):
                form_rule_rows.setdefault(key, []).append(r)

    semantic_by_key: dict[str, str] = {}
    for dossier in raw_form_dossiers:
        d = _as_dict(dossier)
        project_name = _clean(d.get("project_name"))
        form_name = _clean(d.get("form_name"))
        if not form_name:
            continue
        proc_rows = _lookup_rows(form_proc_rows, project_name, form_name)
        sql_rows = _lookup_rows(form_sql_rows, project_name, form_name)
        table_hints = {
            _clean(t)
            for item in sql_rows
            for t in _as_list(_as_dict(item).get("tables"))
            if _clean(t)
        }
        alias = _semantic_form_alias(
            form_name=form_name,
            purpose=_clean(d.get("purpose")),
            db_tables=table_hints,
            procedures=proc_rows,
            rules=_lookup_rows(form_rule_rows, project_name, form_name),
            controls=[_clean(x) for x in _as_list(d.get("controls")) if _clean(x)],
        )
        semantic = _clean(alias).lower()
        if semantic:
            for key in _form_keys(project_name, form_name):
                semantic_by_key[key] = semantic

    donor_rules_by_semantic: dict[str, list[dict[str, Any]]] = {}
    allowed_semantics = {
        "deposit capture",
        "withdrawal processing",
        "transaction ledger",
        "customer management",
        "account type maintenance",
        "password management",
    }
    for key, rows in form_rule_rows.items():
        semantic = semantic_by_key.get(key, "")
        if semantic not in allowed_semantics or not rows:
            continue
        donor_rules_by_semantic.setdefault(semantic, [])
        for r in rows:
            rid = _clean(_as_dict(r).get("rule_id") or _as_dict(r).get("id"))
            existing = {
                _clean(_as_dict(x).get("rule_id") or _as_dict(x).get("id"))
                for x in donor_rules_by_semantic[semantic]
            }
            if rid and rid in existing:
                continue
            donor_rules_by_semantic[semantic].append(r)

    for key, semantic in semantic_by_key.items():
        if semantic not in allowed_semantics:
            continue
        if form_rule_rows.get(key):
            continue
        mirrored = donor_rules_by_semantic.get(semantic, [])
        if mirrored:
            form_rule_rows[key] = mirrored[:8]

    form_control_type_by_key: dict[str, dict[str, str]] = {}
    for row in raw_form_dossiers:
        r = _as_dict(row)
        keys = _form_keys(r.get("project_name"), r.get("form_name"))
        if not keys:
            continue
        for key in keys:
            mapping = form_control_type_by_key.setdefault(key, {})
            for ctl in _as_list(r.get("controls")):
                ctl_text = _clean(ctl)
                if not ctl_text:
                    continue
                if ":" in ctl_text:
                    ctl_type, ctl_name = ctl_text.split(":", 1)
                else:
                    ctl_type, ctl_name = ctl_text, ctl_text
                mapping[_clean(ctl_name).lower()] = _clean(ctl_type)

    shared_module_procs = {
        _clean(_as_dict(proc).get("procedure_name"))
        for proc in raw_procedures
        if _base_form_name(_as_dict(proc).get("form")) == "shared_module"
    }
    for key, events in form_event_rows.items():
        for e in events:
            ev = _as_dict(e)
            for call in [_clean(c) for c in _as_list(ev.get("calls")) if _clean(c)]:
                if call in shared_module_procs:
                    form_shared_components.setdefault(key, set()).add(call)

    dossier_by_key: dict[str, dict[str, Any]] = {}
    for dossier in raw_form_dossiers:
        d = _as_dict(dossier)
        for key in _form_keys(d.get("project_name"), d.get("form_name")):
            if key and key not in dossier_by_key:
                dossier_by_key[key] = d

    discovered_forms: list[dict[str, str]] = []
    seen_discovered: set[str] = set()

    def _normalize_discovered_form_name(value: Any) -> str:
        raw_name = _clean(value)
        if not raw_name:
            return ""
        leaf = Path(raw_name).name
        if ":" in leaf:
            left, right = leaf.split(":", 1)
            if _clean(left).lower() in {"form", "mdiform"} and _clean(right):
                leaf = _clean(right)
        return _clean(leaf)

    def _add_discovered(project_name: Any, form_name: Any, source: str) -> None:
        pname = _clean(project_name)
        fname = _normalize_discovered_form_name(form_name)
        if not fname:
            return
        key = _form_key(pname, fname)
        if not key or key in seen_discovered:
            return
        seen_discovered.add(key)
        discovered_forms.append({"project_name": pname, "form_name": fname, "form_key": key, "source": source})

    for proj in projects:
        p = _as_dict(proj)
        pname = _clean(p.get("name"))
        for form_name in _as_list(p.get("forms")):
            _add_discovered(pname, form_name, "project.forms")
        for asset in _as_list(p.get("ui_assets")):
            a = _as_dict(asset)
            if _clean(a.get("kind")).lower() in {"form", "screen"}:
                _add_discovered(pname, a.get("name"), "project.ui_assets")
        for member in _as_list(p.get("members")):
            m = _as_dict(member)
            kind = _clean(m.get("kind")).lower()
            path = _clean(m.get("path"))
            if kind == "form" or path.lower().endswith(".frm"):
                _add_discovered(pname, Path(path).name, "project.members")

    for dossier in raw_form_dossiers:
        d = _as_dict(dossier)
        _add_discovered(d.get("project_name"), d.get("form_name"), "form_dossier")

    orphan_by_key: dict[str, dict[str, Any]] = {}
    orphan_unmapped_count = 0
    for orphan in raw_orphans:
        o = _as_dict(orphan)
        orphan_form = _clean(o.get("form")) or Path(_clean(o.get("path"))).stem
        orphan_project = _clean(o.get("project_name"))
        if orphan_form == "(unmapped_form_files)":
            summary_text = _clean(o.get("behavior_summary"))
            m = re.search(r"(\\d+)\\s+discovered\\s+form\\s+files", summary_text, flags=re.IGNORECASE)
            if m:
                orphan_unmapped_count = int(m.group(1))
            continue
        key = _form_key(orphan_project, orphan_form)
        if key:
            orphan_by_key[key] = o
            _add_discovered(orphan_project, orphan_form, "orphan_analysis")

    # Section O
    doc.add_page_break()
    _add_paragraph_with_role(doc, "Section O - Project Dependency Map", role="heading1")
    _add_section_intro(doc, _as_dict(plan.get("section_intros")).get("dependency_map", ""))
    project_table = _build_table(doc, ["Project", "Path", "Forms", "Dependencies", "DB touchpoints"])
    for idx, row in enumerate(raw_landscape[:120]):
        r = _as_dict(row)
        raw_id = _clean(r.get("id"))
        left = raw_id.split("|", 1)[0] if "|" in raw_id else raw_id
        path = _clean(r.get("path"))
        project_name = left or "variant"
        if project_display_name_counts.get(project_name, 0) > 1 and path:
            project_name = f"{project_name} ({path})"
        _add_table_row(
            project_table,
            [
                project_name,
                path or "n/a",
                str(_to_int(_as_dict(r.get("counts")).get("forms"))),
                ", ".join(_as_list(r.get("dependencies"))[:6]) or "n/a",
                ", ".join(_as_list(r.get("db_touchpoints"))[:8]) or "n/a",
            ],
            alt=(idx % 2 == 1),
        )

    dep_rows: list[dict[str, str]] = []
    dep_seen: set[tuple[str, str, str, str]] = set()
    for entry in raw_event_map:
        e = _as_dict(entry)
        source = _clean(e.get("container") or e.get("form") or e.get("name")) or "n/a"
        trigger_control = _clean(_as_dict(e.get("trigger")).get("control")).lower()
        for call in [_clean(x) for x in _as_list(e.get("calls")) if _clean(x)]:
            dep_type = ""
            call_norm = _clean(call)
            call_low = call_norm.lower()
            if call in shared_module_procs:
                dep_type = "shared_module_call"
            elif (
                "main" in source.lower()
                or "toolbar" in source.lower()
                or "toolbar" in trigger_control
            ) and call.lower().startswith(("frm", "form", "rpt", "datareport")):
                if call_low.startswith(("rpt", "datareport")):
                    dep_type = "report_navigation"
                elif call_low in {"frm", "form"}:
                    dep_type = "mdi_navigation_unresolved"
                    call_norm = f"{call_norm} [Unresolved]"
                else:
                    dep_type = "mdi_navigation"
            if not dep_type:
                continue
            stable_evidence = _clean(_as_dict(e.get("handler")).get("symbol")) or f"{source}->{call}"
            key = (source, call, dep_type, stable_evidence)
            if key in dep_seen:
                continue
            dep_seen.add(key)
            blocks = "Sprint 1"
            if dep_type == "report_navigation":
                blocks = "Sprint 2"
            elif dep_type == "mdi_navigation_unresolved":
                blocks = "n/a (unresolved)"
            elif dep_type == "cross_variant_schema_conflict":
                blocks = "Sprint 0"
            dep_rows.append(
                {
                    "from": source,
                    "to": call_norm,
                    "type": dep_type,
                    "evidence": stable_evidence,
                    "blocks_sprint": blocks,
                }
            )

    schema = _as_dict(raw_variant_diff.get("schema_divergence"))
    for pair in _as_list(schema.get("blocking_pairs") or schema.get("pairs")):
        p = _as_dict(pair)
        left = _clean(p.get("left_project"))
        right = _clean(p.get("right_project"))
        if not left or not right:
            continue
        evidence = (
            f"alias_mismatches={len(_as_list(p.get('alias_mismatches')))}, "
            f"near_miss={len(_as_list(p.get('near_miss_names')))}, "
            f"transaction_conflict={'yes' if bool(p.get('transaction_schema_conflict')) else 'no'}"
        )
        key = (left, right, "cross_variant_schema_conflict", evidence)
        if key in dep_seen:
            continue
        dep_seen.add(key)
        dep_rows.append(
            {
                "from": left,
                "to": right,
                "type": "cross_variant_schema_conflict",
                "evidence": evidence,
                "blocks_sprint": "Sprint 0",
            }
        )

    dep_table = _build_table(doc, ["From", "To", "Type", "Evidence", "Blocks Sprint"])
    if not dep_rows:
        dep_rows = [
            {"from": "n/a", "to": "n/a", "type": "n/a", "evidence": "No project dependencies detected", "blocks_sprint": "n/a"}
        ]
    for idx, row in enumerate(dep_rows[:600]):
        _add_table_row(
            dep_table,
            [row.get("from"), row.get("to"), row.get("type"), row.get("evidence"), row.get("blocks_sprint")],
            alt=(idx % 2 == 1),
            size=8,
        )

    # Section K
    doc.add_page_break()
    _add_paragraph_with_role(doc, "Section K - Form Dossiers (Extended)", role="heading1")
    _add_section_intro(doc, _as_dict(plan.get("section_intros")).get("form_dossiers", ""))
    form_table = _build_table(
        doc,
        [
            "Form",
            "Display Name",
            "Project",
            "form_type",
            "status",
            "Purpose",
            "Inputs (data)",
            "Outputs (effects)",
            "ActiveX used",
            "DB tables",
            "actions",
            "coverage_score",
            "confidence_score",
            "exclusion_reason",
        ],
    )
    excluded_rows: list[dict[str, str]] = []
    for idx, form_ref in enumerate(
        sorted(discovered_forms, key=lambda x: (_clean(x.get("project_name")), _clean(x.get("form_name")).lower()))[:900]
    ):
        project_name = _clean(form_ref.get("project_name"))
        form_name = _clean(form_ref.get("form_name"))
        form_key = _clean(form_ref.get("form_key")) or _form_key(project_name, form_name)
        base_key = _base_only_key(form_name)
        dossier = _as_dict(dossier_by_key.get(form_key) or dossier_by_key.get(base_key))
        orphan_row = _as_dict(orphan_by_key.get(form_key) or orphan_by_key.get(base_key))

        status = "mapped" if dossier else "excluded"
        exclusion_reason = ""
        if status != "mapped":
            exclusion_reason = _clean(orphan_row.get("recommendation")) or "missing_from_form_dossier"

        purpose = _clean(dossier.get("purpose") or orphan_row.get("behavior_summary"))
        proc_rows_for_form = _lookup_rows(form_proc_rows, project_name, form_name)
        sql_rows_for_form = _lookup_rows(form_sql_rows, project_name, form_name)
        form_rules = _lookup_rows(form_rule_rows, project_name, form_name)
        db_tables_set = _lookup_set(form_db_tables, project_name, form_name)
        db_tables = sorted(db_tables_set)
        form_controls = _as_list(dossier.get("controls"))

        alias = _semantic_form_alias(
            form_name=form_name,
            purpose=purpose,
            db_tables=db_tables_set,
            procedures=proc_rows_for_form,
            rules=form_rules,
            controls=[_clean(x) for x in form_controls if _clean(x)],
        )
        display_name = _display_form_name(form_name, alias) or form_name
        form_type = _infer_form_type(
            form_name=form_name,
            purpose=purpose,
            procedures=proc_rows_for_form,
            controls=form_controls,
            tables=db_tables_set,
        )

        input_values: set[str] = set()
        for ctl in form_controls:
            ctl_text = _clean(ctl)
            if not ctl_text:
                continue
            ctl_name = _clean(ctl_text.split(":", 1)[-1])
            if _is_data_input_control(ctl_name):
                input_values.add(_to_business_input(ctl_name))
        for proc in proc_rows_for_form:
            p = _as_dict(proc)
            for raw_input in _as_list(p.get("inputs")):
                token = _clean(raw_input).split(".", 1)[0]
                if _is_data_input_control(token):
                    input_values.add(_to_business_input(token))

        output_values: set[str] = set()
        for sql_row in sql_rows_for_form:
            sr = _as_dict(sql_row)
            op = _clean(sr.get("operation") or sr.get("kind"))
            for table in _as_list(sr.get("tables")):
                effect = _business_effect_from_sql(op, _clean(table))
                if effect:
                    output_values.add(effect)
        if not output_values:
            fallback = _fallback_business_effects(
                alias=alias,
                purpose=purpose,
                inputs=input_values,
                db_tables=db_tables,
                rules=form_rules,
            )
            output_values.update(fallback)

        form_activex: set[str] = set()
        for ctl in form_controls:
            ctl_text = _clean(ctl)
            if not ctl_text:
                continue
            ctl_type = _clean(ctl_text.split(":", 1)[0])
            if ctl_type and not ctl_type.upper().startswith("VB"):
                form_activex.add(ctl_type)
        for dep in project_dependencies_by_name.get(project_name, set()):
            dep_name = _clean(dep)
            if dep_name.lower().endswith((".ocx", ".dll")) or "MSCOM" in dep_name.upper() or "MSFLEX" in dep_name.upper():
                form_activex.add(dep_name)

        coverage = float(_as_dict(dossier.get("coverage")).get("coverage_score") or 0)
        raw_conf = float(_as_dict(dossier.get("coverage")).get("confidence_score") or 0)
        action_count = len(_as_list(dossier.get("actions"))) if dossier else len(proc_rows_for_form)
        generic_purpose = _clean(purpose).lower() in {
            "business workflow executed through event-driven ui controls.",
            "business workflow executed through event-driven ui controls",
            "potential orphan flow detected.",
            "potential orphan flow detected",
        }
        coverage_clamped = max(0.0, min(1.0, coverage))
        confidence = 0.22 + (0.45 * coverage_clamped)
        confidence += min(0.14, 0.02 * action_count)
        confidence += min(0.08, 0.015 * len(proc_rows_for_form))
        confidence += min(0.08, 0.02 * len(db_tables))
        confidence += 0.09 if sql_rows_for_form else -0.08
        confidence += 0.08 if not generic_purpose else -0.12
        if not input_values:
            confidence -= 0.06
        if action_count == 0:
            confidence -= 0.16
        if bool(re.fullmatch(r"(form\d+|frm\d+)", form_name.lower())) and not _clean(alias):
            confidence -= 0.08
        if 0 < raw_conf <= 1 and abs(raw_conf - 0.92) > 1e-4:
            confidence = (confidence * 0.8) + (raw_conf * 0.2)
        confidence = max(0.1, min(0.98, confidence))

        _add_table_row(
            form_table,
            [
                _qualified_form_name(project_name, form_name),
                display_name or "n/a",
                _project_label(project_name, project_path_by_name),
                form_type,
                status,
                purpose or "n/a",
                ", ".join(sorted(input_values)[:8]) or "n/a",
                ", ".join(sorted(output_values)[:8]) or "n/a",
                ", ".join(sorted(form_activex)[:6]) or "n/a",
                ", ".join(db_tables[:8]) or "n/a",
                str(action_count),
                f"{coverage:.2f}",
                f"{confidence:.2f}",
                exclusion_reason or "n/a",
            ],
            alt=(idx % 2 == 1),
            size=8,
        )

        if status != "mapped":
            excluded_rows.append(
                {
                    "form": _qualified_form_name(project_name, form_name),
                    "reason": exclusion_reason or "missing_from_form_dossier",
                    "source": _clean(form_ref.get("source")) or "detected",
                }
            )

    if excluded_rows or orphan_unmapped_count > 0:
        _add_paragraph_with_role(doc, "K1. Excluded/Unresolved Forms", role="heading2")
        excluded_table = _build_table(doc, ["Form", "Reason", "Source"])
        for idx, row in enumerate(excluded_rows[:400]):
            _add_table_row(
                excluded_table,
                [row.get("form"), row.get("reason"), row.get("source")],
                alt=(idx % 2 == 1),
                size=8,
            )
        if orphan_unmapped_count > 0:
            _add_table_row(
                excluded_table,
                [
                    "(unmapped_form_files)",
                    f"reconcile_project_membership ({orphan_unmapped_count} unresolved form files)",
                    "orphan_analysis",
                ],
                alt=(len(excluded_rows) % 2 == 1),
                size=8,
            )

    _add_paragraph_with_role(doc, "Business Rules by Form", role="heading2")
    rule_table = _build_table(doc, ["Rule ID", "Form", "Layer", "Category", "Business Meaning", "Risk links"])
    rules_rendered = 0
    seen_rule_form: set[tuple[str, str]] = set()

    def _rule_layer(category: str, evidence_text: str, statement_text: str) -> str:
        layer = "Presentation"
        low = f"{evidence_text} {statement_text}".lower()
        if category.lower() in {"data_persistence", "calculation_logic", "threshold_rule"} or any(
            x in low for x in ["select ", "insert ", "update ", "delete ", "table"]
        ):
            layer = "Data"
        if any(x in low for x in [".bas", "module", "shared"]):
            layer = "Shared"
        return layer

    for row in raw_rules[:1200]:
        r = _as_dict(row)
        statement = _clean(r.get("statement"))
        evidence = ", ".join(
            [
                _clean(_as_dict(ev).get("external_ref", {}).get("ref") or _as_dict(ev).get("file_span", {}).get("path"))
                for ev in _as_list(r.get("evidence"))
            ]
        )
        rule_forms: list[str] = []
        for source in [statement, evidence]:
            for form_name in _extract_forms_from_text(source):
                if form_name not in rule_forms:
                    rule_forms.append(form_name)
        if not rule_forms:
            rule_forms = [_clean(_as_dict(r.get("scope")).get("component_id")) or "n/a"]
        category = _clean(r.get("category") or r.get("rule_type") or "other")
        meaning = _rule_business_meaning(statement, category)
        if ("splash" in " ".join(rule_forms).lower() or "splash" in evidence.lower()) and "balance is recalculated" in meaning.lower():
            meaning = "Splash/loading behavior advances progress state before opening workflow screens."
        layer = _rule_layer(category, evidence, statement)
        rule_id = _clean(r.get("rule_id") or r.get("id")) or "BR"

        for form_value in rule_forms[:8]:
            base_form = _base_form_name(form_value) or _base_form_name(_clean(_as_dict(r.get("scope")).get("component_id")))
            pair = (rule_id, base_form or "n/a")
            if pair in seen_rule_form:
                continue
            seen_rule_form.add(pair)
            related_risk_ids: set[str] = set()
            if base_form:
                for dossier in raw_form_dossiers:
                    d = _as_dict(dossier)
                    if _base_form_name(d.get("form_name")) != base_form:
                        continue
                    for risk_row in _lookup_rows(form_risk_rows, d.get("project_name"), d.get("form_name")):
                        rid = _clean(_as_dict(risk_row).get("risk_id"))
                        if rid:
                            related_risk_ids.add(rid)
            rule_low = statement.lower()
            for risk_row in raw_risks:
                rr = _as_dict(risk_row)
                rid = _clean(rr.get("risk_id"))
                desc = _clean(rr.get("description")).lower()
                if not rid:
                    continue
                if any(token in desc and token in rule_low for token in ["caption", "balance", "customerid", "delete", "injection", "credential", "password"]):
                    related_risk_ids.add(rid)
            _add_table_row(
                rule_table,
                [
                    rule_id,
                    form_value or "n/a",
                    layer,
                    category,
                    meaning or statement or "n/a",
                    ", ".join(sorted(related_risk_ids)[:6]) or "none",
                ],
                alt=(rules_rendered % 2 == 1),
                size=8,
            )
            rules_rendered += 1

    for dossier in raw_form_dossiers:
        d = _as_dict(dossier)
        project_name = _clean(d.get("project_name"))
        form_name = _clean(d.get("form_name"))
        base_form = _base_form_name(form_name)
        if not base_form:
            continue
        mirrored_rows = _lookup_rows(form_rule_rows, project_name, form_name)
        for mr in mirrored_rows[:8]:
            rule_id = _clean(_as_dict(mr).get("rule_id") or _as_dict(mr).get("id")) or "BR"
            pair = (rule_id, base_form)
            if pair in seen_rule_form:
                continue
            seen_rule_form.add(pair)
            category = _clean(_as_dict(mr).get("category") or _as_dict(mr).get("rule_type") or "other")
            meaning = _rule_business_meaning(_clean(_as_dict(mr).get("statement")), category)
            _add_table_row(
                rule_table,
                [rule_id, _qualified_form_name(project_name, form_name), "Data", category, meaning or "n/a", "none"],
                alt=(rules_rendered % 2 == 1),
                size=8,
            )
            rules_rendered += 1

    if rules_rendered == 0:
        _add_table_row(rule_table, ["BR-NA", "n/a", "n/a", "n/a", "No business rules available", "none"], alt=False)

    # Section P
    doc.add_page_break()
    _add_paragraph_with_role(doc, "Section P - Form Flow Traces", role="heading1")
    _add_section_intro(doc, _as_dict(plan.get("section_intros")).get("flow_traces", ""))
    rendered_trace_forms: set[str] = set()
    for form_idx, row in enumerate(raw_form_dossiers[:120], start=1):
        r = _as_dict(row)
        form_name = _clean(r.get("form_name")) or f"Form-{form_idx}"
        project_name = _clean(r.get("project_name"))
        trace_form_key = _form_key(project_name, form_name)
        if trace_form_key and trace_form_key in rendered_trace_forms:
            continue
        if trace_form_key:
            rendered_trace_forms.add(trace_form_key)
        _add_paragraph_with_role(doc, f"{_qualified_form_name(project_name, form_name)} ({_project_label(project_name, project_path_by_name)})", role="heading2")

        events = _lookup_rows(form_event_rows, project_name, form_name)
        procedures = _lookup_rows(form_proc_rows, project_name, form_name)
        sql_entries = _lookup_rows(form_sql_rows, project_name, form_name)
        control_map = _lookup_control_map(form_control_type_by_key, project_name, form_name)

        procedure_names: set[str] = set()
        for proc in procedures:
            name = _clean(_as_dict(proc).get("procedure_name"))
            if name:
                procedure_names.add(name)
        for sql_entry in sql_entries:
            name = _clean(_as_dict(sql_entry).get("procedure"))
            if name:
                procedure_names.add(name)

        trace_table = _build_table(doc, ["Callable", "Kind", "Event", "ActiveX", "SQL IDs", "Tables", "Trace status"])
        if not procedure_names:
            _add_table_row(
                trace_table,
                [
                    "n/a",
                    "n/a",
                    "n/a",
                    "n/a",
                    "n/a",
                    ", ".join(sorted(_lookup_set(form_db_tables, project_name, form_name))[:8]) or "n/a",
                    "TRACE_GAP",
                ],
                alt=False,
                size=8,
            )
            _set_cell_shading(trace_table.rows[1].cells[6], COLOR_RISK_MED_BG)
            continue

        for idx, proc_name in enumerate(sorted(procedure_names)[:120]):
            related_events = [_as_dict(e) for e in events if proc_name in _clean(_as_dict(e).get("handler", {}).get("symbol"))]
            related_sql = [_as_dict(s) for s in sql_entries if _clean(_as_dict(s).get("procedure")) == proc_name]
            activex_hits: list[str] = []
            for e in related_events:
                trigger_control = _clean(_as_dict(e.get("trigger")).get("control"))
                ctl_type = _clean(control_map.get(trigger_control.lower()))
                if trigger_control and ctl_type and not ctl_type.upper().startswith("VB"):
                    activex_hits.append(f"{trigger_control}:{ctl_type}")
            sql_ids = sorted({_clean(x.get("sql_id")) for x in related_sql if _clean(x.get("sql_id"))})
            table_names = sorted({_clean(t) for x in related_sql for t in _as_list(_as_dict(x).get("tables")) if _clean(t)})
            event_refs = set(event_handler_by_key_proc.get((_form_key(project_name, form_name), proc_name), set()))
            if not event_refs:
                event_refs = {_clean(_as_dict(e.get("handler")).get("symbol")) for e in related_events if _clean(_as_dict(e.get("handler")).get("symbol"))}
            trace_ok = bool(sql_ids) and bool(table_names)
            kind = _callable_kind(
                proc_name,
                form_name,
                _clean(_as_dict(related_events[0].get("trigger")).get("event")) if related_events else "",
            )
            _add_table_row(
                trace_table,
                [
                    proc_name,
                    kind,
                    ", ".join(sorted(event_refs)[:3]) or "n/a",
                    ", ".join(sorted(set(activex_hits))[:5]) or "n/a",
                    ", ".join(sql_ids[:6]) or "n/a",
                    ", ".join(table_names[:8]) or "n/a",
                    "OK" if trace_ok else "TRACE_GAP",
                ],
                alt=(idx % 2 == 1),
                size=8,
            )
            _set_cell_shading(trace_table.rows[idx + 1].cells[6], COLOR_RISK_LOW_BG if trace_ok else COLOR_RISK_MED_BG)

    # Section Q
    doc.add_page_break()
    _add_paragraph_with_role(doc, "Section Q - Form Traceability Matrix", role="heading1")
    _add_section_intro(doc, _as_dict(plan.get("section_intros")).get("traceability", ""))
    trace_table = _build_table(
        doc,
        [
            "Form",
            "Project",
            "has_event_map",
            "has_sql_map",
            "has_business_rules",
            "has_risk_entry",
            "completeness_score",
            "missing_links",
        ],
    )
    traceability_rows: list[dict[str, Any]] = []
    seen_trace_keys: set[str] = set()
    for idx, form_ref in enumerate(
        sorted(discovered_forms, key=lambda x: (_clean(x.get("project_name")), _clean(x.get("form_name")).lower()))[:500]
    ):
        form_name = _clean(form_ref.get("form_name")) or "n/a"
        project_name = _clean(form_ref.get("project_name"))
        form_key = _clean(form_ref.get("form_key")) or _form_key(project_name, form_name)
        if form_key in seen_trace_keys:
            continue
        seen_trace_keys.add(form_key)
        has_event = bool(_lookup_rows(form_event_rows, project_name, form_name))
        has_sql = bool(_lookup_rows(form_sql_rows, project_name, form_name))
        has_rules = bool(_lookup_rows(form_rule_rows, project_name, form_name))
        has_risk = bool(_lookup_rows(form_risk_rows, project_name, form_name))
        has_proc = bool(_lookup_rows(form_proc_rows, project_name, form_name))
        missing: list[str] = []
        if not has_event:
            missing.append("event_map")
        if not has_sql:
            missing.append("sql_map")
        if not has_rules:
            missing.append("business_rules")
        if not has_risk:
            missing.append("risk_register")
        if not has_proc:
            missing.append("procedure_summary")
        completeness_score = int((int(has_event) + int(has_sql) + int(has_rules) + int(has_risk) + int(has_proc)) * 20)
        _add_table_row(
            trace_table,
            [
                _qualified_form_name(project_name, form_name),
                _project_label(project_name, project_path_by_name),
                "yes" if has_event else "no",
                "yes" if has_sql else "no",
                "yes" if has_rules else "no",
                "yes" if has_risk else "no",
                str(completeness_score),
                ", ".join(missing) or "none",
            ],
            alt=(idx % 2 == 1),
            size=8,
        )
        traceability_rows.append(
            {
                "form_name": form_name,
                "project_name": project_name,
                "form_key": form_key,
                "qualified_form": _qualified_form_name(project_name, form_name),
                "missing": missing,
                "risk_ids": [
                    _clean(_as_dict(rr).get("risk_id"))
                    for rr in _lookup_rows(form_risk_rows, project_name, form_name)[:4]
                    if _clean(_as_dict(rr).get("risk_id"))
                ],
            }
        )

    # Section R
    _add_paragraph_with_role(doc, "Section R - Sprint Dependency Map", role="heading1")
    _add_section_intro(doc, _as_dict(plan.get("section_intros")).get("sprints", ""))
    sprint_table = _build_table(doc, ["Form", "Suggested sprint", "Depends on", "Shared Components Required", "Rationale"])
    variant_gate_needed = bool(raw_variant_diff.get("decision_required"))
    sorted_rows = sorted(
        traceability_rows,
        key=lambda row: (len(_as_list(row.get("missing"))), 0 if _as_list(row.get("risk_ids")) else 1, _clean(row.get("qualified_form"))),
        reverse=True,
    )
    if not sorted_rows:
        sorted_rows = [{"form_name": "n/a", "project_name": "", "form_key": "", "qualified_form": "n/a", "missing": ["event_map"], "risk_ids": []}]
    emitted: set[str] = set()
    for idx, row in enumerate(sorted_rows[:500]):
        form_key = _clean(row.get("form_key"))
        if form_key and form_key in emitted:
            continue
        if form_key:
            emitted.add(form_key)
        missing = _as_list(row.get("missing"))
        risk_ids = _as_list(row.get("risk_ids"))
        deps: list[str] = []
        if variant_gate_needed:
            deps.append("DEC-VARIANT-001")
        if "sql_map" in missing:
            deps.append("Q.sql_map")
        if "event_map" in missing:
            deps.append("Q.event_map")
        if "business_rules" in missing:
            deps.append("Q.business_rules")
        if risk_ids:
            deps.extend([_clean(x) for x in risk_ids[:2] if _clean(x)])
        if "event_map" in missing or "sql_map" in missing:
            sprint = "Sprint 0 (Discovery closure)"
            rationale = "Close traceability gaps before modernization changes."
        elif risk_ids:
            sprint = "Sprint 1 (Risk-first modernization)"
            rationale = "Implement remediation-first changes for high-risk legacy behavior."
        else:
            sprint = "Sprint 2 (Parity hardening)"
            rationale = "Complete hardening, regression validation, and release evidence for production readiness."
        shared_required = sorted(_lookup_set(form_shared_components, row.get("project_name"), row.get("form_name")))
        _add_table_row(
            sprint_table,
            [
                _clean(row.get("qualified_form")) or _clean(row.get("form_name")) or "n/a",
                sprint,
                ", ".join(deps) or "none",
                ", ".join(shared_required[:5]) or "none",
                rationale,
            ],
            alt=(idx % 2 == 1),
            size=8,
        )

    # Risks and references
    _add_paragraph_with_role(doc, "Risks Register Snapshot", role="heading1")
    _add_section_intro(doc, _as_dict(plan.get("section_intros")).get("risks", ""))
    risk_table = _build_table(doc, ["Risk ID", "Severity", "Description", "Recommended action"])
    risks_to_render = raw_risks[:220] if raw_risks else _as_list(data.get("top_risks"))
    if not risks_to_render:
        risks_to_render = [{"risk_id": "RISK-NA", "severity": "low", "description": "No explicit risks listed.", "recommended_action": ""}]
    for idx, row in enumerate(risks_to_render):
        r = _as_dict(row)
        sev = _clean(r.get("severity")) or "medium"
        _add_table_row(
            risk_table,
            [
                _clean(r.get("risk_id") or r.get("id")) or "RISK",
                sev.upper(),
                _clean(r.get("description")) or "n/a",
                _clean(r.get("recommended_action") or r.get("mitigation")) or "n/a",
            ],
            alt=(idx % 2 == 1),
            size=8,
        )
        _set_cell_shading(risk_table.rows[idx + 1].cells[1], _severity_fill(sev))

    _add_paragraph_with_role(doc, "Evidence References", role="heading2")
    refs = _as_dict(data.get("artifact_refs"))
    ref_table = _build_table(doc, ["Artifact", "Reference"])
    if not refs:
        _add_table_row(ref_table, ["n/a", "No artifact references listed."], alt=False)
    for idx, (key, value) in enumerate(list(refs.items())[:120]):
        _add_table_row(ref_table, [key, _clean(value) or "n/a"], alt=(idx % 2 == 1), size=8)

    disclaimer = doc.add_paragraph()
    _set_run_style(
        disclaimer.add_run(
            "Generated from analyst structured artifacts. This workbook is intended for BA/business review "
            "and should be paired with technical JSON/markdown exports for engineering implementation."
        ),
        size=8,
        color=COLOR_MUTED,
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _xml_text(value: str) -> str:
    return html.escape(str(value or ""), quote=False)


def _paragraph_xml(text: str, *, bold: bool = False) -> str:
    if not text:
        return "<w:p/>"
    escaped = _xml_text(text)
    if bold:
        return (
            "<w:p><w:r><w:rPr><w:b/></w:rPr>"
            f"<w:t xml:space=\"preserve\">{escaped}</w:t>"
            "</w:r></w:p>"
        )
    return "<w:p><w:r>" f"<w:t xml:space=\"preserve\">{escaped}</w:t>" "</w:r></w:p>"


def _document_xml(paragraphs: list[tuple[str, bool]]) -> str:
    body = "".join([_paragraph_xml(text, bold=bold) for text, bold in paragraphs])
    return (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document "
        "xmlns:wpc=\"http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas\" "
        "xmlns:mc=\"http://schemas.openxmlformats.org/markup-compatibility/2006\" "
        "xmlns:o=\"urn:schemas-microsoft-com:office:office\" "
        "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\" "
        "xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\" "
        "xmlns:v=\"urn:schemas-microsoft-com:vml\" "
        "xmlns:wp14=\"http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing\" "
        "xmlns:wp=\"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing\" "
        "xmlns:w10=\"urn:schemas-microsoft-com:office:word\" "
        "xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
        "xmlns:w14=\"http://schemas.microsoft.com/office/word/2010/wordml\" "
        "xmlns:wpg=\"http://schemas.microsoft.com/office/word/2010/wordprocessingGroup\" "
        "xmlns:wpi=\"http://schemas.microsoft.com/office/word/2010/wordprocessingInk\" "
        "xmlns:wne=\"http://schemas.microsoft.com/office/word/2006/wordml\" "
        "xmlns:wps=\"http://schemas.microsoft.com/office/word/2010/wordprocessingShape\" "
        "mc:Ignorable=\"w14 wp14\">"
        f"<w:body>{body}<w:sectPr><w:pgSz w:w=\"12240\" w:h=\"15840\"/>"
        "<w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" w:header=\"708\" w:footer=\"708\" w:gutter=\"0\"/>"
        "<w:cols w:space=\"708\"/><w:docGrid w:linePitch=\"360\"/></w:sectPr></w:body></w:document>"
    )


def _content_types_xml() -> str:
    return (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
        "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
        "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
        "<Override PartName=\"/word/document.xml\" "
        "ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
        "<Override PartName=\"/docProps/core.xml\" "
        "ContentType=\"application/vnd.openxmlformats-package.core-properties+xml\"/>"
        "<Override PartName=\"/docProps/app.xml\" "
        "ContentType=\"application/vnd.openxmlformats-officedocument.extended-properties+xml\"/>"
        "</Types>"
    )


def _package_rels_xml() -> str:
    return (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
        "<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>"
        "<Relationship Id=\"rId2\" Type=\"http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties\" Target=\"docProps/core.xml\"/>"
        "<Relationship Id=\"rId3\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties\" Target=\"docProps/app.xml\"/>"
        "</Relationships>"
    )


def _core_props_xml() -> str:
    stamp = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
    return (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<cp:coreProperties xmlns:cp=\"http://schemas.openxmlformats.org/package/2006/metadata/core-properties\" "
        "xmlns:dc=\"http://purl.org/dc/elements/1.1/\" xmlns:dcterms=\"http://purl.org/dc/terms/\" "
        "xmlns:dcmitype=\"http://purl.org/dc/dcmitype/\" xmlns:xsi=\"http://www.w3.org/2001/XMLSchema-instance\">"
        "<dc:title>Synthetix Analyst Business Brief</dc:title>"
        "<dc:creator>Synthetix Analyst Agent</dc:creator>"
        "<cp:lastModifiedBy>Synthetix Analyst Agent</cp:lastModifiedBy>"
        f"<dcterms:created xsi:type=\"dcterms:W3CDTF\">{stamp}</dcterms:created>"
        f"<dcterms:modified xsi:type=\"dcterms:W3CDTF\">{stamp}</dcterms:modified>"
        "</cp:coreProperties>"
    )


def _app_props_xml() -> str:
    return (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<Properties xmlns=\"http://schemas.openxmlformats.org/officeDocument/2006/extended-properties\" "
        "xmlns:vt=\"http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes\">"
        "<Application>Synthetix</Application>"
        "</Properties>"
    )


def _build_business_paragraphs(payload: dict[str, Any], run_id: str) -> list[tuple[str, bool]]:
    report, raw_artifacts = _extract_report_and_raw(payload)
    data = _collect_report_data(report, raw_artifacts)
    lines: list[tuple[str, bool]] = []
    lines.append(("Modernization Business Brief", True))
    lines.append((f"Project: {_clean(data['project'].get('name')) or 'Untitled project'}", False))
    if run_id:
        lines.append((f"Run ID: {run_id}", False))
    lines.append((f"Generated: {_clean(data['metadata'].get('generated_at')) or datetime.now(timezone.utc).isoformat()}", False))
    lines.append((f"Source: {_clean(data['context'].get('repo')) or 'n/a'} @ {_clean(data['context'].get('branch')) or 'main'}", False))
    lines.append(("", False))
    lines.append(("Executive Summary", True))
    objective = _clean(data["project"].get("objective"))
    if objective:
        lines.append((f"Objective: {objective}", False))
    lines.append((f"Readiness Score: {_clean(data['glance'].get('readiness_score')) or 'n/a'}/100", False))
    lines.append((f"Risk Tier: {_clean(data['glance'].get('risk_tier')) or 'n/a'}", False))
    qa_summary = _as_dict(data.get("qa_summary"))
    if qa_summary:
        lines.append(
            (
                "QA Summary: "
                f"status={_clean(qa_summary.get('status')) or 'NOT_RUN'}, "
                f"pass={_to_int(qa_summary.get('pass_count'))}, "
                f"warn={_to_int(qa_summary.get('warn_count'))}, "
                f"fail={_to_int(qa_summary.get('fail_count'))}, "
                f"blockers={_to_int(qa_summary.get('blocker_count'))}",
                False,
            )
        )
    lines.append(("Generated by Synthetix Analyst Agent", False))
    return lines


def _build_business_docx_bytes_fallback(payload: dict[str, Any], *, run_id: str = "") -> bytes:
    paragraphs = _build_business_paragraphs(_as_dict(payload), run_id)
    package = io.BytesIO()
    with ZipFile(package, mode="w", compression=ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _content_types_xml())
        zf.writestr("_rels/.rels", _package_rels_xml())
        zf.writestr("docProps/core.xml", _core_props_xml())
        zf.writestr("docProps/app.xml", _app_props_xml())
        zf.writestr("word/document.xml", _document_xml(paragraphs))
    return package.getvalue()


def build_business_docx_bytes(
    payload: dict[str, Any],
    *,
    run_id: str = "",
    render_mode: str = "deterministic",
    llm_doc_plan: dict[str, Any] | None = None,
    template_path: str | None = None,
    strict_template: bool = False,
) -> bytes:
    safe = _as_dict(payload)
    try:
        use_plan = llm_doc_plan if str(render_mode).strip().lower() == "llm_rich" else None
        return _build_business_docx_bytes_rich(
            safe,
            run_id=run_id,
            llm_doc_plan=use_plan,
            template_path=template_path,
            strict_template=bool(strict_template),
        )
    except Exception:
        return _build_business_docx_bytes_fallback(safe, run_id=run_id)
