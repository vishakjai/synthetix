from __future__ import annotations

import io
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
except Exception:  # pragma: no cover
    Document = None  # type: ignore
    WD_TABLE_ALIGNMENT = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    OxmlElement = None  # type: ignore
    qn = None  # type: ignore
    Inches = None  # type: ignore
    Pt = None  # type: ignore
    RGBColor = None  # type: ignore

try:
    from PIL import Image, ImageDraw, ImageFont  # type: ignore
except Exception:  # pragma: no cover
    Image = None  # type: ignore
    ImageDraw = None  # type: ignore
    ImageFont = None  # type: ignore


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _clean(value: Any) -> str:
    return str(value or "").strip()


def _hex_rgb(value: str) -> Any:
    text = _clean(value).replace("#", "") or "111111"
    if RGBColor is None:
        return None
    return RGBColor.from_string(text[:6].ljust(6, "0"))


def _resolve_template_path(doc_type: str, template_path: str | None = None) -> str:
    requested = _clean(template_path)
    if requested and Path(requested).exists():
        return requested
    candidates = []
    if _clean(doc_type).lower() == "legacy":
        candidates.append(Path("/Users/vishak/Downloads/Legacy-HLD-Template.docx"))
    else:
        candidates.append(Path("/Users/vishak/Downloads/Target-HLD-Template.docx"))
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return ""


def _clear_document_body(doc: Any) -> None:
    body = getattr(getattr(doc, "_element", None), "body", None)
    if body is None:
        return
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _set_run_style(run: Any, *, size: int = 10, bold: bool = False, italic: bool = False, color: str = "111111") -> None:
    font = getattr(run, "font", None)
    if font is None:
        return
    if Pt is not None:
        font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    rgb = _hex_rgb(color)
    if rgb is not None:
        font.color.rgb = rgb


def _set_cell_shading(cell: Any, color: str) -> None:
    if OxmlElement is None or qn is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _clean(color).replace("#", ""))
    tc_pr.append(shd)


def _set_cell_text(cell: Any, text: str, *, color: str = "111111", bold: bool = False, italic: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(_clean(text))
    _set_run_style(run, size=size, bold=bold, italic=italic, color=color)


def _build_table(doc: Any, columns: list[str], rows: list[dict[str, Any]], *, brand_color: str) -> Any:
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER if WD_TABLE_ALIGNMENT else 1
    table.autofit = False
    for idx, heading in enumerate(columns):
        cell = table.rows[0].cells[idx]
        _set_cell_shading(cell, brand_color)
        _set_cell_text(cell, heading, color="FFFFFF", bold=True, size=9)
    alt = False
    for row in rows:
        cells = table.add_row().cells
        for idx, heading in enumerate(columns):
            normalized_heading = heading.lower().replace(" ", "_").replace("/", "_").replace("-", "_")
            candidate_keys = {
                normalized_heading,
                normalized_heading.replace("mapped_", ""),
                normalized_heading.replace("legacy_", ""),
                normalized_heading.replace("target_", ""),
                normalized_heading.replace("refs", "ref"),
            }
            value = ""
            for key, candidate in row.items():
                if _clean(key).lower().replace(" ", "_").replace("/", "_").replace("-", "_") in candidate_keys:
                    value = _clean(candidate)
                    break
            _set_cell_text(
                cells[idx],
                value,
                color="E8511A" if value.startswith("[") else "111111",
                italic=value.startswith("["),
                size=9,
            )
            if alt:
                _set_cell_shading(cells[idx], "F5F8FB")
        alt = not alt
    return table


def _add_toc(doc: Any) -> None:
    if OxmlElement is None:
        return
    p = doc.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_separate)
    txt = p.add_run("Update field in Word to render the table of contents.")
    _set_run_style(txt, size=9, italic=True, color="E8511A")
    r._r.append(fld_end)


def _add_heading(doc: Any, text: str, *, level: int = 1, color: str = "111111") -> None:
    style_name = f"Heading {level}" if level in {1, 2, 3} else None
    try:
        paragraph = doc.add_paragraph(style=style_name)
    except Exception:
        paragraph = doc.add_paragraph()
    run = paragraph.add_run(_clean(text))
    _set_run_style(run, size=14 - level, bold=True, color=color)


def _diagram_png_bytes(title: str, mermaid: str, brand_color: str) -> bytes | None:
    if Image is None or ImageDraw is None:
        return None
    width = 1400
    height = 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default() if ImageFont else None
    brand = f"#{_clean(brand_color).replace('#', '')[:6]}"
    draw.rectangle((40, 40, width - 40, height - 40), outline=brand, width=4)
    draw.text((70, 70), _clean(title), fill=brand, font=font)
    lines = [_clean(line) for line in _clean(mermaid).splitlines() if _clean(line)]
    y = 130
    for line in lines[:24]:
        draw.text((70, y), line[:120], fill="#111111", font=font)
        y += 28
    out = io.BytesIO()
    image.save(out, format="PNG", dpi=(150, 150))
    return out.getvalue()


def build_architect_hld_docx_bytes(
    payload: dict[str, Any],
    *,
    template_path: str | None = None,
    strict_template: bool = True,
) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not available")
    safe = _as_dict(payload)
    doc_type = _clean(safe.get("document_type")) or "legacy"
    resolved_template = _resolve_template_path(doc_type, template_path) if strict_template else _clean(template_path)
    doc = Document(resolved_template) if resolved_template else Document()
    if resolved_template:
        _clear_document_body(doc)

    section = doc.sections[0]
    if Inches is not None:
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)

    brand_color = _clean(safe.get("brand_color")).replace("#", "") or "1B3A5C"
    accent_color = _clean(safe.get("accent_color")).replace("#", "") or "E8511A"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH else 1
    title_run = title.add_run(_clean(safe.get("title")))
    _set_run_style(title_run, size=18, bold=True, color=brand_color)

    meta_table = _build_table(
        doc,
        ["Field", "Value"],
        [{"field": key, "value": value} for key, value in _as_list(safe.get("metadata_rows"))],
        brand_color=brand_color,
    )
    if meta_table.rows:
        for row in meta_table.rows[1:]:
            _set_cell_text(row.cells[0], _clean(row.cells[0].text), bold=True, size=9)

    warning_table = _build_table(
        doc,
        ["Human Review Warnings"],
        [{"human_review_warnings": value} for value in _as_list(safe.get("cover_warnings"))],
        brand_color=accent_color,
    )
    for row in warning_table.rows[1:]:
        _set_cell_shading(row.cells[0], "FFF4E5")

    doc.add_page_break()
    _add_heading(doc, "Table of Contents", level=1, color=brand_color)
    _add_toc(doc)
    doc.add_page_break()

    figure_index = 1
    for section_payload in _as_list(safe.get("sections")):
        if not isinstance(section_payload, dict):
            continue
        _add_heading(doc, _clean(section_payload.get("title")), level=1, color=brand_color)
        for paragraph_text in _as_list(section_payload.get("paragraphs")):
            para = doc.add_paragraph()
            run = para.add_run(_clean(paragraph_text))
            _set_run_style(run, size=10, color=accent_color if _clean(paragraph_text).startswith("[") else "111111", italic=_clean(paragraph_text).startswith("["))
        for bullet_text in _as_list(section_payload.get("bullets")):
            try:
                para = doc.add_paragraph(style="List Bullet")
            except Exception:
                para = doc.add_paragraph()
            run = para.add_run(_clean(bullet_text))
            _set_run_style(run, size=9, color=accent_color if _clean(bullet_text).startswith("[") else "111111", italic=_clean(bullet_text).startswith("["))
        if bool(section_payload.get("figure")):
            png_bytes = _diagram_png_bytes(_clean(safe.get("diagram_title")), _clean(safe.get("diagram_mermaid")), brand_color)
            if png_bytes is not None:
                picture_stream = io.BytesIO(png_bytes)
                if Inches is not None:
                    doc.add_picture(picture_stream, width=Inches(6.7))
                else:
                    doc.add_picture(picture_stream)
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH else 1
                cap_run = caption.add_run(f"Figure {figure_index}: {_clean(safe.get('diagram_title'))} - Generated by Synthetix Architect Agent")
                _set_run_style(cap_run, size=9, italic=True, color="6B7280")
                figure_index += 1
            else:
                para = doc.add_paragraph()
                run = para.add_run("[ Populate diagram figure from generated Mermaid output before client distribution. Source: architect diagrams. ]")
                _set_run_style(run, size=9, italic=True, color=accent_color)
        for table_payload in _as_list(section_payload.get("tables")):
            if not isinstance(table_payload, dict):
                continue
            if _clean(table_payload.get("title")):
                _add_heading(doc, _clean(table_payload.get("title")), level=2, color=brand_color)
            _build_table(
                doc,
                [str(value) for value in _as_list(table_payload.get("columns"))],
                [row for row in _as_list(table_payload.get("rows")) if isinstance(row, dict)],
                brand_color=brand_color,
            )

    doc.add_page_break()
    _add_heading(doc, "Appendix - Artefact Reference Hashes", level=1, color=brand_color)
    appendix_rows = [{"artifact": key, "sha256": value} for key, value in sorted(_as_dict(safe.get("appendix_hashes")).items())]
    _build_table(doc, ["Artifact", "SHA256"], appendix_rows, brand_color=brand_color)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
